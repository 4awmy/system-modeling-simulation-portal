import docx
path = 'c:/Users/omarh/OneDrive/Desktop/Uni/System Modeling and Simulation/Lectures/Curriclum.docx'
doc = docx.Document(path)
with open('scripts/curriculum_extracted.txt', 'w', encoding='utf-8') as f:
    f.write(f"Total Paragraphs: {len(doc.paragraphs)}\n\n")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            f.write(f"[{i}]: {p.text}\n")
print("Successfully extracted curriculum text!")
